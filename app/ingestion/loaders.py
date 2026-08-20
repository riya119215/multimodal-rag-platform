import os
import json
from pathlib import Path
from typing import List, Optional, Union
import pypdf
import docx
import pandas as pd

from app.ingestion.metadata import Document, DocumentChunk
from app.utils.helpers import clean_text, format_timestamp
from app.core.logging_config import logger

class BaseLoader:
    """Base interface for all document loaders."""
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        raise NotImplementedError

class TextLoader(BaseLoader):
    """Loads plain text or markdown files."""
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        path = Path(file_path)
        logger.info(f"Loading text file: {path.name}")
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = clean_text(f.read())
        
        doc_type = "markdown" if path.suffix.lower() == ".md" else "text"
        return [
            Document(
                content=content,
                source_file=path.name,
                doc_type=doc_type,
                metadata={"file_size": path.stat().st_size, "path": str(path)}
            )
        ]

class PDFLoader(BaseLoader):
    """Loads PDF documents page-by-page preserving page numbers."""
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        path = Path(file_path)
        logger.info(f"Loading PDF document: {path.name}")
        reader = pypdf.PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages, start=1):
            text = clean_text(page.extract_text() or "")
            if not text:
                continue
            docs.append(
                Document(
                    content=text,
                    source_file=path.name,
                    doc_type="pdf",
                    metadata={
                        "page_number": i,
                        "total_pages": len(reader.pages),
                        "path": str(path)
                    }
                )
            )
        return docs

class DocxLoader(BaseLoader):
    """Loads Word DOCX documents."""
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        path = Path(file_path)
        logger.info(f"Loading DOCX document: {path.name}")
        doc = docx.Document(str(path))
        paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
        
        # Include tables if any
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(clean_text(cell.text) for cell in row.cells if clean_text(cell.text))
                if row_text:
                    table_texts.append(row_text)

        full_content = "\n\n".join(paragraphs + table_texts)
        return [
            Document(
                content=full_content,
                source_file=path.name,
                doc_type="docx",
                metadata={"paragraphs": len(paragraphs), "tables": len(doc.tables), "path": str(path)}
            )
        ]

class CSVLoader(BaseLoader):
    """Loads structured CSV files and converts rows into natural text format."""
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        path = Path(file_path)
        logger.info(f"Loading CSV dataset: {path.name}")
        df = pd.read_csv(str(path))
        rows = []
        for idx, row in df.iterrows():
            row_items = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
            rows.append(f"Row {idx + 1}: " + ", ".join(row_items))
        
        content = "\n".join(rows)
        return [
            Document(
                content=content,
                source_file=path.name,
                doc_type="csv",
                metadata={"total_rows": len(df), "columns": list(df.columns), "path": str(path)}
            )
        ]

class AudioTranscriptJSONLoader(BaseLoader):
    """Loads pre-transcribed JSON transcript files with timestamped segments."""
    def load(self, file_path: Union[str, Path]) -> List[DocumentChunk]:
        path = Path(file_path)
        logger.info(f"Loading transcript JSON: {path.name}")
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        source_file = data.get("file_name", path.name)
        video_num = str(data.get("video_id", "N/A"))
        title = data.get("title", path.stem)
        chunks = data.get("chunks", [])

        chunk_objs = []
        for idx, c in enumerate(chunks):
            text = clean_text(c.get("text", ""))
            if not text:
                continue
            start_sec = float(c.get("start", 0.0))
            end_sec = float(c.get("end", 0.0))
            chunk_objs.append(
                DocumentChunk(
                    chunk_id=idx,
                    text=text,
                    source_file=source_file,
                    doc_type="audio_transcript",
                    doc_id=f"audio_{path.stem}_{idx}",
                    video_number=str(c.get("number", video_num)),
                    title=c.get("title", title),
                    start=start_sec,
                    end=end_sec,
                    start_formatted=c.get("start_formatted", format_timestamp(start_sec)),
                    end_formatted=c.get("end_formatted", format_timestamp(end_sec)),
                    metadata={"audio_source": source_file}
                )
            )
        return chunk_objs

class DocumentLoaderFactory:
    """Dispatches appropriate loader based on file extension."""
    @staticmethod
    def get_loader(file_path: Union[str, Path]):
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext in [".txt", ".md"]:
            return TextLoader()
        elif ext == ".pdf":
            return PDFLoader()
        elif ext == ".docx":
            return DocxLoader()
        elif ext == ".csv":
            return CSVLoader()
        elif ext == ".json":
            return AudioTranscriptJSONLoader()
        else:
            raise ValueError(f"Unsupported document format '{ext}' for file: {path.name}")

    @staticmethod
    def load_document(file_path: Union[str, Path]) -> Union[List[Document], List[DocumentChunk]]:
        loader = DocumentLoaderFactory.get_loader(file_path)
        return loader.load(file_path)
